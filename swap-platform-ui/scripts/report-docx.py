#!/usr/bin/env python3
"""
Swap Platform DOCX Report Generator
Usage: python scripts/report-docx.py [reportPath] [outputPath]
  reportPath  default: scripts/last-report.json
  outputPath  default: scripts/last-report.docx
"""

import json
import sys
import os
from pathlib import Path

import json
import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def load_report(path):
    with open(path, encoding='utf-8') as f:
        return json.load(f)

def generate_docx(report, output_path):
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)

    # Title
    title = doc.add_heading('Swap Platform API 测试报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary table
    summary = report.get('summary', {})
    pass_rate = (summary.get('pass', 0) / summary.get('total', 1) * 100) if summary.get('total', 0) > 0 else 0
    icon = '🟢' if summary.get('fail', 0) == 0 else ('🟡' if summary.get('fail', 0) <= summary.get('total', 1) / 2 else '🔴')

    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    cells = table.rows[0].cells
    cells[0].text = '时间'
    cells[1].text = report.get('time', 'N/A')
    cells = table.rows[1].cells
    cells[0].text = '目标'
    cells[1].text = report.get('baseUrl', 'N/A')
    cells = table.rows[2].cells
    cells[0].text = '状态'
    cells[1].text = f"{icon} {summary.get('pass', 0)}/{summary.get('total', 0)} 通过 ({pass_rate:.1f}%)"
    cells = table.rows[3].cells
    cells[0].text = '结果'
    cells[1].text = f"通过: {summary.get('pass', 0)}, 失败: {summary.get('fail', 0)}"

    # Classification stats
    doc.add_paragraph()
    doc.add_heading('分类统计', level=1)
    groups = report.get('groups', {})
    if groups:
        gtable = doc.add_table(rows=len(groups) + 1, cols=2)
        gtable.style = 'Light Grid Accent 1'
        gtable.rows[0].cells[0].text = '模块'
        gtable.rows[0].cells[1].text = '结果'
        for i, (group, result) in enumerate(groups.items(), 1):
            parts = result.split('/')
            p, t = int(parts[0]), int(parts[1])
            emoji = '✅' if p == t else ('❌' if p == 0 else '⚠️')
            gtable.rows[i].cells[0].text = group
            gtable.rows[i].cells[1].text = f"{emoji} {p}/{t} 通过"

    # Failures
    failed = report.get('failed', [])
    if failed:
        doc.add_paragraph()
        doc.add_heading(f'失败用例 ({len(failed)})', level=1)
        for f in failed:
            doc.add_heading(f.get('name', 'Unknown'), level=2)
            p = doc.add_paragraph()
            p.add_run(f'HTTP 状态: {f.get("status", "N/A")}\n')
            p.add_run(f'响应 code: {f.get("code", "N/A")}\n')
            p.add_run(f'响应 msg: {f.get("msg", "N/A")}\n')
            p.add_run(f'原始返回: {f.get("raw", "N/A")}\n')

    # Differences
    diff = report.get('diff', [])
    if diff:
        doc.add_paragraph()
        doc.add_heading('契约差异', level=1)
        for d in diff:
            doc.add_paragraph(d, style='List Bullet')

    # Uncovered
    uncovered = report.get('uncovered', [])
    if uncovered:
        doc.add_paragraph()
        doc.add_heading('未覆盖接口', level=1)
        for u in uncovered:
            doc.add_paragraph(f'{u.get("method", "")} {u.get("path", "")} - {u.get("summary", "")}', style='List Bullet')

    # Passed
    passed = report.get('passed', [])
    if passed:
        doc.add_paragraph()
        doc.add_heading(f'通过用例 ({len(passed)})', level=1)
        for name in passed:
            doc.add_paragraph(name, style='List Bullet')

    doc.save(output_path)
    print(f"DOCX report: {output_path}")


if __name__ == '__main__':
    report_path = sys.argv[1] if len(sys.argv) > 1 else 'last-report.json'
    output_path = sys.argv[2] if len(sys.argv) > 2 else 'last-report.docx'

    # Ensure absolute paths from script location
    script_dir = Path(__file__).parent
    if not os.path.isabs(report_path):
        report_path = script_dir / report_path
    if not os.path.isabs(output_path):
        output_path = script_dir / output_path

    report = load_report(str(report_path))
    generate_docx(report, str(output_path))
